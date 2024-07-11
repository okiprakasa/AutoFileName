from docx.api import Document
import os
from tkinter import filedialog
from tkinter import *

root = Tk()
root.withdraw()
source_folder = filedialog.askdirectory()

# source_folder = r'C:\Users\okipr\Documents\Projects\AutoFileName'
count = 1
fileName = ""
for item in os.listdir(source_folder):
    # check if an item is file or not
    if os.path.isfile(os.path.join(source_folder, item)):
        if item.endswith('.docx'):
            try:
                # Open the .docx file
                doc = Document(item)

                # Create an empty list to store all paragraphs
                paraText = ""
                tableText = ""

                # Iterate through each paragraph and append its text to the list
                for p in doc.paragraphs:
                    paraText = paraText + p.text

                for table in doc.tables:
                    for row in table.rows:
                        tableText = tableText + " ".join([cell.text for cell in row.cells])

                # remove forbidden ASCII characters for file name
                fileName = ((tableText + paraText).replace("\n", " ")
                            .replace("/", "").replace("\\", "")
                            .replace("<", "").replace(">", "")
                            .replace(":", "").replace(";", "")
                            .replace("|", "").replace("?", "")
                            .replace("&", "").replace("!", "")
                            .replace("@", "").replace("#", "")
                            .replace("$", "").replace("^", "")
                            .replace("+", "").replace("-", "")
                            .replace("~", "").replace("`", "")
                            .replace(".", "").replace("*", ""))

                # remove double space
                fileName = ' '.join(fileName.split())

                if len(fileName) > 250:
                    fileName = fileName[:250]

                os.rename(
                    os.path.join(source_folder, item),
                    os.path.join(source_folder, fileName + '.docx')
                )
            except PermissionError:
                continue
            except FileExistsError:
                fileName = fileName + str(count)
                os.rename(
                    os.path.join(source_folder, item),
                    os.path.join(source_folder, fileName + '.docx')
                )
                count += 1
            except Exception as e:
                print(e)
                continue
